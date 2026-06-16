"""Document-to-OpenClaw memory helpers."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
import re
import shutil
import subprocess
from typing import Any

import yaml

from .memory_captions import _frontmatter_value, _relative_posix, resolve_memory_root


SUPPORTED_DOCUMENT_EXTENSIONS = {".pdf", ".doc", ".docx", ".epub", ".txt", ".md"}

_SPACE_RE = re.compile(r"[ \t]+")
_BLANK_LINES_RE = re.compile(r"\n{3,}")
_SOURCE_ID_RE = re.compile(r"[^a-zA-Z0-9._-]+")


@dataclass(frozen=True)
class DocumentText:
    text: str
    metadata: dict[str, Any]


@dataclass(frozen=True)
class DocumentImportResult:
    markdown_path: Path
    parsed_text_path: Path
    raw_source_path: Path | None
    word_count: int
    character_count: int


def discover_document_files(input_path: Path) -> list[Path]:
    if input_path.is_file():
        return [input_path]
    if not input_path.is_dir():
        raise FileNotFoundError(f"document input not found: {input_path}")
    return sorted(
        path
        for path in input_path.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_DOCUMENT_EXTENSIONS
    )


def infer_document_id(path: Path) -> str:
    candidate = _SOURCE_ID_RE.sub("-", path.stem.strip()).strip("-._")
    return candidate or "document"


def _clean_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = "\n".join(_SPACE_RE.sub(" ", line).strip() for line in normalized.splitlines())
    return _BLANK_LINES_RE.sub("\n\n", normalized).strip()


def _extract_pdf_text(path: Path) -> DocumentText:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError(
            "pypdf is required for PDF imports. Install with: pip install pypdf"
        ) from exc

    reader = PdfReader(str(path))
    page_text = [page.extract_text() or "" for page in reader.pages]
    return DocumentText(
        text=_clean_text("\n\n".join(page_text)),
        metadata={"page_count": len(reader.pages)},
    )


def _extract_docx_text(path: Path) -> DocumentText:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required for DOCX imports. Install with: pip install python-docx"
        ) from exc

    document = Document(str(path))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return DocumentText(text=_clean_text("\n\n".join(parts)), metadata={})


def _extract_doc_text(path: Path) -> DocumentText:
    try:
        completed = subprocess.run(
            ["antiword", str(path)],
            check=True,
            capture_output=True,
            text=True,
        )
    except FileNotFoundError as exc:
        raise RuntimeError(
            "Legacy .doc imports require antiword on PATH, or convert the file to .docx/PDF first."
        ) from exc
    except subprocess.CalledProcessError as exc:
        raise RuntimeError(f"antiword failed to extract {path}: {exc.stderr.strip()}") from exc
    return DocumentText(text=_clean_text(completed.stdout), metadata={})


def _extract_epub_text(path: Path) -> DocumentText:
    try:
        from bs4 import BeautifulSoup
        from ebooklib import ITEM_DOCUMENT, epub
    except ImportError as exc:
        raise RuntimeError(
            "ebooklib and beautifulsoup4 are required for EPUB imports. "
            "Install with: pip install ebooklib beautifulsoup4"
        ) from exc

    book = epub.read_epub(str(path))
    parts: list[str] = []
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        soup = BeautifulSoup(item.get_content(), "html.parser")
        text = soup.get_text("\n")
        if text.strip():
            parts.append(text)
    metadata: dict[str, Any] = {}
    creators = [value for value, _attrs in book.get_metadata("DC", "creator")]
    if creators:
        metadata["authors"] = creators
    return DocumentText(text=_clean_text("\n\n".join(parts)), metadata=metadata)


def extract_document_text(path: Path) -> DocumentText:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_text(path)
    if suffix == ".docx":
        return _extract_docx_text(path)
    if suffix == ".doc":
        return _extract_doc_text(path)
    if suffix == ".epub":
        return _extract_epub_text(path)
    if suffix in {".txt", ".md"}:
        return DocumentText(
            text=_clean_text(path.read_text(encoding="utf-8", errors="replace")), metadata={}
        )
    raise ValueError(f"unsupported document type: {path.suffix}")


def _copy_raw_file(source: Path, raw_dir: Path) -> Path:
    raw_dir.mkdir(parents=True, exist_ok=True)
    destination = raw_dir / source.name
    if source.resolve() != destination.resolve():
        shutil.copy2(source, destination)
    return destination


def _write_parsed_text(path: Path, text: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(f"{text}\n" if text else "", encoding="utf-8")
    return path


def _word_count(text: str) -> int:
    return len(re.findall(r"\S+", text))


def _markdown_text_sections(text: str, *, chunk_size: int = 8000) -> list[str]:
    if not text:
        return ["No extractable text was found."]
    if len(text) <= chunk_size:
        return [text]

    sections: list[str] = []
    offset = 0
    section_number = 1
    while offset < len(text):
        end = min(len(text), offset + chunk_size)
        split_at = text.rfind("\n\n", offset, end)
        if split_at <= offset:
            split_at = end
        chunk = text[offset:split_at].strip()
        if chunk:
            sections.extend([f"### Section {section_number}", "", chunk])
            section_number += 1
        offset = split_at
    return sections


def build_document_markdown(
    *,
    source_id: str,
    title: str,
    source_type: str,
    source_url: str | None,
    author: str | None,
    raw_source_rel: str | None,
    parsed_text_rel: str,
    imported_at: str,
    document_text: DocumentText,
    enrichment: dict | None = None,
) -> str:
    text = document_text.text
    word_count = _word_count(text)
    frontmatter = {
        "source_type": source_type,
        "source_id": source_id,
        "title": title,
        "author": _frontmatter_value(author),
        "source_url": _frontmatter_value(source_url),
        "source_file": _frontmatter_value(raw_source_rel),
        "parsed_text_file": parsed_text_rel,
        "imported_at": imported_at,
        "word_count": word_count,
        "character_count": len(text),
    }
    for key, value in document_text.metadata.items():
        frontmatter[key] = _frontmatter_value(value)

    enrichment = enrichment or {}
    keywords = enrichment.get("keywords") or []
    cleaned_keywords = [str(keyword).strip() for keyword in keywords if str(keyword).strip()]
    if cleaned_keywords:
        frontmatter["keywords"] = cleaned_keywords
    frontmatter = {key: value for key, value in frontmatter.items() if value is not None}

    body: list[str] = ["---"]
    body.append(yaml.safe_dump(frontmatter, sort_keys=False, allow_unicode=True).strip())
    body.append("---")
    body.extend(["", f"# {title}", ""])
    if source_url:
        body.append(f"Source: {source_url}  ")
    if author:
        body.append(f"Author: {author}  ")
    body.append(f"Source type: {source_type}  ")
    body.append(f"Words: {word_count}  ")
    body.append("")

    body.extend(["## Source Files", ""])
    body.append(f"- Parsed text: `{parsed_text_rel}`")
    if raw_source_rel:
        body.append(f"- Original source: `{raw_source_rel}`")
    if source_url:
        body.append(f"- Source URL: {source_url}")
    body.append("")

    short_summary = str(enrichment.get("short_summary") or "").strip()
    long_description = str(enrichment.get("long_description") or "").strip()
    if short_summary:
        body.extend(["## Short Summary", "", short_summary, ""])
    if long_description:
        body.extend(["## Longer Description", "", long_description, ""])
    if cleaned_keywords:
        body.extend(
            ["## Keywords", "", ", ".join(f"`{keyword}`" for keyword in cleaned_keywords), ""]
        )

    body.extend(["## Parsed Text", ""])
    body.extend(_markdown_text_sections(text))
    body.append("")
    return "\n".join(body)


def import_document_file(
    *,
    document_file: Path,
    collection: str,
    workspace: Path | None = None,
    memory_dir: Path | None = None,
    source_id: str | None = None,
    title: str | None = None,
    source_url: str | None = None,
    source_type: str | None = None,
    author: str | None = None,
    copy_raw: bool = True,
    enrichment: dict | None = None,
    imported_at: str | None = None,
) -> DocumentImportResult:
    document_file = document_file.resolve()
    memory_root = resolve_memory_root(workspace=workspace, memory_dir=memory_dir)
    collection = collection.strip().strip("/")
    if not collection:
        raise ValueError("collection is required")

    resolved_source_id = (source_id or infer_document_id(document_file)).strip()
    resolved_title = (title or document_file.stem).strip()
    resolved_type = (source_type or document_file.suffix.lower().lstrip(".") or "document").strip()
    resolved_imported_at = imported_at or datetime.now(timezone.utc).isoformat()

    collection_dir = memory_root / collection
    documents_dir = collection_dir / "documents"
    raw_dir = collection_dir / "raw"
    documents_dir.mkdir(parents=True, exist_ok=True)

    document_text = extract_document_text(document_file)
    parsed_text_path = _write_parsed_text(
        collection_dir / "parsed" / f"{resolved_source_id}.txt",
        document_text.text,
    )
    raw_source_path = _copy_raw_file(document_file, raw_dir) if copy_raw else None

    parsed_text_rel = _relative_posix(parsed_text_path, documents_dir)
    raw_source_rel = _relative_posix(raw_source_path, documents_dir) if raw_source_path else None
    markdown = build_document_markdown(
        source_id=resolved_source_id,
        title=resolved_title,
        source_type=resolved_type,
        source_url=source_url,
        author=author,
        raw_source_rel=raw_source_rel,
        parsed_text_rel=parsed_text_rel,
        imported_at=resolved_imported_at,
        document_text=document_text,
        enrichment=enrichment,
    )

    markdown_path = documents_dir / f"{resolved_source_id}.md"
    markdown_path.write_text(markdown, encoding="utf-8")
    return DocumentImportResult(
        markdown_path=markdown_path,
        parsed_text_path=parsed_text_path,
        raw_source_path=raw_source_path,
        word_count=_word_count(document_text.text),
        character_count=len(document_text.text),
    )
