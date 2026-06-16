from pathlib import Path

import pytest
import yaml
from typer.testing import CliRunner

from hypercli_cli.cli import app
from hypercli_cli.memory_documents import extract_document_text, import_document_file


runner = CliRunner()


def _write(path: Path, text: str) -> Path:
    path.write_text(text, encoding="utf-8")
    return path


def _write_docx(path: Path) -> Path:
    from docx import Document

    document = Document()
    document.add_heading("Creator Brief", level=1)
    document.add_paragraph("DOCX memory text about a branded creator agent.")
    document.add_paragraph("The voice, examples, and source citations matter.")
    document.save(path)
    return path


def _write_epub(path: Path) -> Path:
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_identifier("memory-epub")
    book.set_title("Memory EPUB")
    book.add_author("Example Author")
    chapter = epub.EpubHtml(title="Chapter 1", file_name="chap_01.xhtml", lang="en")
    chapter.content = (
        "<html><body><h1>Chapter 1</h1>"
        "<p>EPUB memory text about channel lore and recurring topics.</p></body></html>"
    )
    book.add_item(chapter)
    book.spine = ["nav", chapter]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())
    epub.write_epub(str(path), book)
    return path


def _write_pdf(path: Path) -> Path:
    from pypdf import PdfWriter
    from pypdf.generic import DictionaryObject, NameObject, StreamObject

    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
    )
    content = StreamObject()
    content._data = b"BT /F1 12 Tf 72 720 Td (PDF memory text for OpenClaw ingestion.) Tj ET"
    page[NameObject("/Contents")] = writer._add_object(content)
    with path.open("wb") as handle:
        writer.write(handle)
    return path


@pytest.mark.parametrize(
    ("factory", "needle"),
    [
        (lambda path: _write(path, "Plain text memory source."), "Plain text memory source."),
        (_write_docx, "DOCX memory text about a branded creator agent."),
        (_write_epub, "EPUB memory text about channel lore"),
        (_write_pdf, "PDF memory text for OpenClaw ingestion."),
    ],
)
def test_extract_document_text_supported_formats(tmp_path, factory, needle):
    suffix = {
        "Plain text memory source.": ".txt",
        "DOCX memory text about a branded creator agent.": ".docx",
        "EPUB memory text about channel lore": ".epub",
        "PDF memory text for OpenClaw ingestion.": ".pdf",
    }[needle]
    document = factory(tmp_path / f"source{suffix}")

    extracted = extract_document_text(document)

    assert needle in extracted.text


def test_import_document_file_writes_rich_markdown_parsed_text_and_raw(tmp_path):
    document = _write(
        tmp_path / "creator-notes.txt", "Creator notes\n\nRecurring format: field tests."
    )
    workspace = tmp_path / "workspace"

    result = import_document_file(
        document_file=document,
        workspace=workspace,
        collection="docs/example",
        source_id="creator-notes",
        title="Creator Notes",
        source_url="https://example.test/notes",
        author="Example Creator",
        imported_at="2026-05-29T00:00:00+00:00",
        enrichment={
            "short_summary": "Notes about the creator's recurring format.",
            "long_description": "The document captures source material for a personality memory.",
            "keywords": ["creator notes", "field tests"],
        },
    )

    markdown = result.markdown_path.read_text(encoding="utf-8")
    frontmatter = yaml.safe_load(markdown.split("---", 2)[1])

    assert result.markdown_path == workspace / "memory/docs/example/documents/creator-notes.md"
    assert result.parsed_text_path == workspace / "memory/docs/example/parsed/creator-notes.txt"
    assert result.raw_source_path == workspace / "memory/docs/example/raw/creator-notes.txt"
    assert result.parsed_text_path.read_text(encoding="utf-8") == (
        "Creator notes\n\nRecurring format: field tests.\n"
    )
    assert frontmatter["source_url"] == "https://example.test/notes"
    assert frontmatter["source_file"] == "../raw/creator-notes.txt"
    assert frontmatter["parsed_text_file"] == "../parsed/creator-notes.txt"
    assert frontmatter["keywords"] == ["creator notes", "field tests"]
    assert "- Parsed text: `../parsed/creator-notes.txt`" in markdown
    assert "- Original source: `../raw/creator-notes.txt`" in markdown
    assert "## Parsed Text" in markdown
    assert "Recurring format: field tests." in markdown


def test_memory_import_documents_cli_uses_hyper_memory_dir(monkeypatch, tmp_path):
    document = _write(tmp_path / "notes.txt", "A simple source document for memory.")
    memory_dir = tmp_path / "explicit-memory"
    monkeypatch.setenv("HYPER_MEMORY_DIR", str(memory_dir))

    result = runner.invoke(
        app,
        [
            "memory",
            "import",
            str(document),
            "--collection",
            "docs/example",
            "--title",
            "Source Notes",
        ],
    )

    assert result.exit_code == 0, result.stdout
    markdown_path = memory_dir / "docs/example/documents/notes.md"
    assert markdown_path.exists()
    assert "# Source Notes" in markdown_path.read_text(encoding="utf-8")
    assert (memory_dir / "docs/example/raw/notes.txt").exists()


def test_memory_import_cli_dispatches_mixed_directory_by_extension(tmp_path):
    synthetic_srt = """1
00:00:00,000 --> 00:00:01,000
Hello from captions.
"""
    caption = _write(tmp_path / "video123.en.srt", synthetic_srt)
    document = _write(tmp_path / "notes.txt", "A simple source document for memory.")
    workspace = tmp_path / "workspace"

    result = runner.invoke(
        app,
        [
            "memory",
            "import",
            str(tmp_path),
            "--workspace",
            str(workspace),
            "--collection",
            "mixed/example",
        ],
    )

    assert result.exit_code == 0, result.stdout
    assert (workspace / "memory/mixed/example/videos/video123.md").exists()
    assert (workspace / "memory/mixed/example/documents/notes.md").exists()
    assert (workspace / f"memory/mixed/example/raw/{caption.name}").exists()
    assert (workspace / f"memory/mixed/example/raw/{document.name}").exists()
